import tkinter as tk
import threading
import pyaudio
import wave
import time
import speech_recognition as sr
# docx
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
# 2pdf
from docx2pdf import convert
# Open file
from tkinter import filedialog
# Save as
from tkinter.filedialog import asksaveasfile


class App():
    chunk = 1024 
    sample_format = pyaudio.paInt16 
    channels = 2
    fs = 44100  
    r = sr.Recognizer()
    file_path = ''
    
    frames = []  
    def __init__(self, master):
        self.isrecording = False
        self.isrecording = True
        # nested frames
        self.frame1 = tk.Frame(main)
        self.frame2 = tk.Frame(main)
        self.frame3 = tk.Frame(main)
        self.frame4 = tk.Frame(main)

        self.btn_explorer = tk.Button(self.frame1, text='Guardar como...', command=self.save_file)
        self.button1 = tk.Button(main, text='Grabar', bg='blue', fg='white', command=self.startrecording)
        self.button2 = tk.Button(main, text='Detener', bg='red', fg='black', command=self.stoprecording)
        self.l1 = tk.Label(self.frame1, text="Documento de salida (.docx): ")
        # checks
        self.state_var = tk.IntVar()
        self.audio_state = tk.Checkbutton(self.frame2, text=" Renombrar audio?", variable=self.state_var)
        self.state_var2 = tk.IntVar()
        self.pdf_state = tk.Checkbutton(self.frame2, text=" Generar pdf?", variable=self.state_var2)
        # label and entry title
        self.doc_title = tk.Label(self.frame3, text="Nombre del acta: ")
        self.title_entry = tk.Entry(self.frame3, width=30)
        # subject
        self.l2 = tk.Label(self.frame4, text="Asunto de la reunión: ")
        self.subject_box = tk.Text(self.frame4, height= 5, width=50, wrap="word")

        self.l3 = tk.Label(main, text="Estado...")

        # saveas
        self.l1.pack(side=tk.LEFT)
        self.l1.pack(padx=10, pady=5)
        self.btn_explorer.pack(side=tk.LEFT)
        self.btn_explorer.pack(padx=10, pady=5)
        # checks
        self.audio_state.pack(side=tk.LEFT) 
        self.audio_state.pack(padx=10, pady=5)
        self.pdf_state.pack(side=tk.LEFT)
        self.pdf_state.pack(padx=10, pady=5)
        # doc title
        self.doc_title.pack(side=tk.LEFT)
        self.doc_title.pack(padx=10, pady=5)
        self.title_entry.pack(side=tk.LEFT)
        self.title_entry.pack(padx=10, pady=5)
        # subject
        self.l2.pack()
        self.l2.pack(padx=10, pady=5)
        self.subject_box.pack()
        self.subject_box.pack(padx=10, pady=5)
        # nested frames
        self.frame1.pack(side=tk.TOP)
        self.frame2.pack(side=tk.TOP)
        self.frame3.pack(side=tk.TOP)
        self.frame4.pack()

        self.l3.pack(padx=5, pady=5)

        self.button1.pack(side=tk.LEFT)
        self.button1.pack(padx=70, pady=5)

        self.button2.pack(side=tk.RIGHT)
        self.button2.pack(padx=70, pady=5)


    def startrecording(self):
        self.p = pyaudio.PyAudio()  
        self.stream = self.p.open(
                format=self.sample_format,
                channels=self.channels,
                rate=self.fs,
                frames_per_buffer=self.chunk,
                input=True)

        self.isrecording = True

        # Creando el documento
        self.mydoc = docx.Document()
        if not self.title_entry.get():
            self.mydoc.add_heading("Reunion " + time.strftime("%d-%m-%Y"), 0)
        else:
            self.mydoc.add_heading(self.title_entry.get(), 0)

        self.mydoc.add_paragraph(time.strftime("Fecha: %d/%m/%Y"))
        self.mydoc.add_paragraph(time.strftime("Hora inicio: %H:%M:%S"))
        self.mydoc.add_paragraph("Asunto: " + self.subject_box.get("1.0",
        "end-1c"))

        print('Recording')
        self.l3.configure(text="Grabando...")
        # Creando hilo
        t = threading.Thread(target=self.record)
        t.start()

    def stoprecording(self):
        self.isrecording = False

        print('recording complete')
        self.l3.configure(text="Finalizado.")
        # Quitando el nombre del pdf del path:
        index = self.file_path.rindex('/')
        self.audioname = self.file_path[:index+1]
        if self.state_var.get() is not 0:
            index = self.file_path.rindex('.')
            self.audioname = self.file_path[:index+1] + "wav"
        else:
            self.audioname += time.strftime("reunion%d_%m_%Y_%H%M%S.wav")

        # Guarda el audio de salida:
        wf = wave.open(self.audioname, 'wb')
        wf.setnchannels(self.channels)
        wf.setsampwidth(self.p.get_sample_size(self.sample_format))
        wf.setframerate(self.fs)
        wf.writeframes(b''.join(self.frames))
        wf.close()

        #---- inicio de la supuesta funcion
        recorded_audio = sr.AudioFile(self.audioname)
        with recorded_audio as source:
            audio = self.r.record(source)
            try:
                text = self.r.recognize_google(audio, language='es-ES')
                # Agregando la transcripcion al documento
                self.mydoc.add_heading("Transcripción", 1)
                my_paragraph = self.mydoc.add_paragraph(text)
                my_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self.mydoc.add_paragraph(time.strftime("Hora final: %H:%M:%S"))
                        
                # Guardando el archivo
                self.mydoc.save(self.file_path)
                # Generando el pdf
                if self.state_var2.get() == 1:
                    convert(self.file_path)

            except:
                self.l3.configure(text="Error en el reconocimiento!")
        #---- fin de la supuesta funcion

        #main.destroy()

    def record(self):
        while self.isrecording:
            data = self.stream.read(self.chunk)
            self.frames.append(data)
        if self.isrecording is False:
            # Deteniendo la grabacion:
            self.stream.stop_stream()
            self.stream.close()
            # Finaliza la interfaz PortAudio
            self.p.terminate()

    def browser_files(self):
        filename = filedialog.askopenfilename(
                initialdir = ".",
                title = "Select a file",
                filetypes = (("docx files", ".docx"),
                ("all files", "*.*")))

    def save_file(self):
        global file_path
        files = [("Todos los archivos", "*.*"),
                ("Documento de Word", "*.docx")]
        file = asksaveasfile(filetypes = files, defaultextension=files)
        self.file_path = file.name



main = tk.Tk()
main.title('Speech Translator')
main.geometry('600x320')

app = App(main)
main.mainloop()

